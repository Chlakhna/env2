# import streamlit as st
# import pandas as pd
# import json
# import os
# import zipfile
# from docx import Document
# from docx.shared import Pt
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.enum.section import WD_SECTION
# import smtplib
# from email.mime.multipart import MIMEMultipart
# from email.mime.text import MIMEText
# from email.mime.base import MIMEBase
# from email import encoders
# import openai
# from datetime import datetime
# import re
# # import plotly.express as px
# import requests
# from docx2pdf import convert
# from io import BytesIO
# from dotenv import load_dotenv
# import time
# from fpdf import FPDF


# # Your existing code

# # Load environment variables from .env file
# load_dotenv()

# # Set your OpenAI API key from environment variable
# # openai.api_key = os.getenv('OPENAI_API_KEY')
# openai.api_key = st.secrets["OPENAI_API_KEY"]

# # Set your Telegram bot token and chat ID
# telegram_bot_token = os.getenv('TELEGRAM_BOT_TOKEN')
# telegram_chat_id = '-1002164741954'  # This can stay hardcoded or also be moved to an environment variable if needed

# # Email configuration
# from_email = "seaklav168@gmail.com"
# password = os.getenv('EMAIL_PASSWORD')

# def send_to_telegram(file_path, caption):
#     url = f"https://api.telegram.org/bot{telegram_bot_token}/sendDocument"
#     with open(file_path, 'rb') as file:
#         response = requests.post(url, data={'chat_id': telegram_chat_id, 'caption': caption}, files={'document': file})
#     return response

# def generate_report_with_chatgpt(data, report_title):
#     try:
#         prompt = (
#             f"Please give a formal report based on provided data and contents as below:  "
#             "Insert a page break here.\n\n"
#             "table of content"
#             "Please generate a formal Table of Contents for a report. The Table of Contents should have each section title followed by a series of dots (dot leaders) leading to the corresponding page number. Ensure the page numbers are aligned to the right margin, like in the example below:"
#             "1. Introduction................................. 1"
#             "2. Data Visualization........................... 5"
#             "3. Results...................................... 9"
#             "4. Conclusion and Recommendations............... 12"
#             "The document should be formatted professionally, and the Table of Contents should follow the above style, with consistent dot leaders and right-aligned page numbers."
#             "Insert a page break here.\n\n"
#             "Abstract:"
#             "Insert a page break here.\n\n"
#             "1. Introduction "
#             "Please describe more details about each of the subheadings:"
#             "1.1. Demographic Profile: Describe respondents, age, gender."
#             "1.2. Land Ownership and Cultivation: Describe their overall land."
#             "1.3. Horticulture Practices: Describe crops, land, and yield of each year."
#             "1.4 Satisfaction Rates: The overall satisfaction."
#             "2. Data Visualization "
#             "In sheet, there are huge differences before and after the project. We want to know how much land, frequency, and yield have increased or decreased after the project. Please provide a table and graph comparing the increase or decrease of land, frequency of plant, and the yield."
#             "Also, describe the percentage of crops that have been planted as well."
#             "3. Discussion and Results "
#             "Overall results "
#             "4. Conclusion and Recommendations: conclude everything and give recommendations."
#             "Please describe each part more detailed."
#             f"{json.dumps(data, indent=2)}"
#         )

#         response = openai.ChatCompletion.create(
#             model="gpt-4o-mini",  # Or "gpt-3.5-turbo"
#             messages=[{"role": "system", "content": "You are a helpful assistant."}, {"role": "user", "content": prompt}],
#             max_tokens=8000,
#             temperature=0.7
#         )

#         report_parts = response.choices[0].message['content'].strip()

#         # Store report parts in session state to avoid re-running
#         # st.session_state["report_parts"] = report_parts

#         st.write(report_parts)

#         # Save report as Word and PDF documents with the dynamic title
#         word_filename = f'{report_title}.docx'
#         pdf_filename = f'{report_title}.pdf'
#         save_report_as_word(report_parts, word_filename)
#         convert_to_pdf_with_retry(word_filename, pdf_filename)

#         return report_parts, word_filename, pdf_filename
#     except Exception as e:
#         st.error(f"Failed to generate report: {e}")
#         return None, None, None

# def create_cover_page(doc, report_title):
#     section = doc.sections[0]
#     section.page_height = Pt(842)  # A4 size height
#     section.page_width = Pt(595)   # A4 size width

#     # Set margins to 2 cm (56.7 points)
#     for section in doc.sections:
#         section.top_margin = Pt(56.7)
#         section.bottom_margin = Pt(56.7)
#         section.left_margin = Pt(56.7)
#         section.right_margin = Pt(56.7)

#     # Add logo from URL to the cover page
#     logo_url = "https://dcxsea.com/asset/images/logo/LOGO_DCX.png"
#     response = requests.get(logo_url)
#     if response.status_code == 200:
#         logo_image = BytesIO(response.content)
#         paragraph = doc.add_paragraph()
#         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

#         run = paragraph.add_run()
#         run.add_picture(logo_image, width=Pt(150))  # Adjust width as needed

#         # Add company name after the logo
#         run = paragraph.add_run("  DCx Co., Ltd.")
#         run.font.size = Pt(24)
#         run.bold = True
#     else:
#         st.error("Failed to download the logo image.")

#     doc.add_paragraph("\n")
#     cover = doc.add_paragraph()
#     cover.alignment = WD_ALIGN_PARAGRAPH.CENTER

#     run = cover.add_run("Report")
#     run.font.size = Pt(20)
#     run.bold = True
#     cover.add_run("\n").font.size = Pt(24)

#     # Add some spacing
#     doc.add_paragraph("\n")
#     cover = doc.add_paragraph()
#     cover.alignment = WD_ALIGN_PARAGRAPH.CENTER

#     run = cover.add_run("Indigenous Agriculture Adaptation")
#     run.font.size = Pt(24)
#     run.bold = True
#     cover.add_run("\n").font.size = Pt(24)

#     for _ in range(2):
#         doc.add_paragraph("\n")

#     # Add prepared for
#     prepared_for = doc.add_paragraph()
#     prepared_for.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     run = prepared_for.add_run("Prepared for: Jack Jasmin")
#     run.font.size = Pt(14)
#     run.add_break()  # Add a break line
#     run.add_break()
#     run = prepared_for.add_run("Prepared by: Black Eye Team")
#     run.font.size = Pt(14)
#     run.add_break()
#     run.add_break()

#     # Place the date at the bottom center of the cover page
#     footer = section.footer
#     footer_paragraph = footer.paragraphs[0]
#     footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     run = footer_paragraph.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
#     run.font.size = Pt(14)

#     # Create a new section for the rest of the document, so the footer with the date is not repeated
#     new_section = doc.add_section(WD_SECTION.NEW_PAGE)

#     # Ensure the new section starts with a fresh header and footer
#     new_section.footer.is_linked_to_previous = False
#     new_section.footer.paragraphs[0].clear()
#     new_section.header.is_linked_to_previous = False

# def add_markdown_formatted_text(paragraph, text):
#     bold_italic_pattern = re.compile(r'\*\*\*(.+?)\*\*\*')
#     bold_pattern = re.compile(r'\*\*(.+?)\*\*')
#     italic_pattern = re.compile(r'\*(.+?)\*')
#     patterns = [
#         ('bold_italic', bold_italic_pattern),
#         ('bold', bold_pattern),
#         ('italic', italic_pattern)
#     ]

#     def replace_match(match, style):
#         if style == 'bold_italic':
#             run = paragraph.add_run(match.group(1))
#             run.bold = True
#             run.italic = True
#         elif style == 'bold':
#             run = paragraph.add_run(match.group(1))
#             run.bold = True
#         elif style == 'italic':
#             run = paragraph.add_run(match.group(1))
#             run.italic = True

#     cursor = 0
#     text_length = len(text)
#     while cursor < text_length:
#         nearest_match = None
#         nearest_style = None
#         nearest_start = text_length

#         for style, pattern in patterns:
#             match = pattern.search(text, cursor)
#             if match:
#                 start = match.start()
#                 if start < nearest_start:
#                     nearest_match = match
#                     nearest_style = style
#                     nearest_start = start

#         if nearest_match:
#             if nearest_start > cursor:
#                 paragraph.add_run(text[cursor:nearest_start])
#             replace_match(nearest_match, nearest_style)
#             cursor = nearest_match.end()
#         else:
#             paragraph.add_run(text[cursor:])
#             break

#     paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# def save_report_as_word(report, filename):
#     try:
#         doc = Document()

#         for section in doc.sections:
#             section.top_margin = Pt(56.7)
#             section.bottom_margin = Pt(56.7)
#             section.left_margin = Pt(56.7)
#             section.right_margin = Pt(56.7)

#         create_cover_page(doc, filename.split('.')[0])

#         lines = report.split('\n')
#         table = None
#         lines = lines[1:]

#         for line in lines:
#             if line.strip() == "---":
#                 doc.add_page_break()
#             elif line.strip().startswith("# "):
#                 paragraph = doc.add_heading(line.strip()[2:], level=1)
#             elif line.strip().startswith("## "):
#                 paragraph = doc.add_heading(line.strip()[3:], level=2)
#             elif line.strip().startswith("### "):
#                 paragraph = doc.add_heading(line.strip()[4:], level=3)
#             elif line.strip().startswith("#### "):
#                 paragraph = doc.add_heading(line.strip()[5:], level=4)
#             elif line.strip().startswith("* "):
#                 doc.add_paragraph(line.strip()[2:], style='List Bullet')
#             elif "|" in line:
#                 table_data = [cell.strip() for cell in line.split('|') if cell]
#                 if not table:
#                     table = doc.add_table(rows=1, cols=len(table_data))
#                     hdr_cells = table.rows[0].cells
#                     for i, cell_data in enumerate(table_data):
#                         hdr_cells[i].text = cell_data
#                 else:
#                     row_cells = table.add_row().cells
#                     for i, cell_data in enumerate(table_data):
#                         row_cells[i].text = cell_data
#             else:
#                 paragraph = doc.add_paragraph()
#                 add_markdown_formatted_text(paragraph, line)

#         doc.save(filename)
#     except Exception as e:
#         st.error(f"Failed to save Word report: {e}")

# # def convert_to_pdf_with_retry(word_filename, pdf_filename, retries=3, delay=5):
# #     try:
# #         pythoncom.CoInitialize()
# #         for attempt in range(retries):
# #             try:
# #                 convert(word_filename, pdf_filename)
# #                 return
# #             except Exception as e:
# #                 st.error(f"Attempt {attempt + 1} failed: {e}")
# #                 if attempt < retries - 1:
# #                     time.sleep(delay)
# #                 else:
# #                     st.error("Failed to convert Word to PDF after multiple attempts.")
# #     finally:
# #         pythoncom.CoUninitialize()
# # def convert_to_pdf_with_retry(word_filename, pdf_filename, retries=3, delay=5):
# #     for attempt in range(retries):
# #         try:
# #             convert(word_filename, pdf_filename)
# #             st.success("Conversion successful!")
# #             return
# #         except Exception as e:
# #             st.error(f"Attempt {attempt + 1} failed: {e}")
# #             if attempt < retries - 1:
# #                 time.sleep(delay)
# #             else:
# #                 st.error("Failed to convert Word to PDF after multiple attempts.")

# # def create_zip_file(word_filename, pdf_filename, zip_filename):
# #     try:
# #         with zipfile.ZipFile(zip_filename, 'w') as zipf:
# #             zipf.write(word_filename)
# #             zipf.write(pdf_filename)
# #         st.success(f"Zip file {zip_filename} created successfully.")
# #     except Exception as e:
# #         st.error(f"Failed to create zip file: {e}")

# # Function to read content from DOCX
# def read_docx_content(word_filename):
#     doc = Document(word_filename)
#     content = ""
#     for para in doc.paragraphs:
#         content += para.text + "\n"
#     return content

# # Function to create PDF using fpdf2
# def create_pdf_from_text(pdf_filename, content):
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     for line in content.split("\n"):
#         if len(line.strip()) > 0:
#             pdf.multi_cell(0, 10, line)
#     pdf.output(pdf_filename)

# # Function to convert DOCX to PDF with retry logic
# def convert_to_pdf_with_retry(word_filename, pdf_filename, retries=3, delay=5):
#     for attempt in range(retries):
#         try:
#             content = read_docx_content(word_filename)
#             create_pdf_from_text(pdf_filename, content)
#             if os.path.exists(pdf_filename):
#                 st.success("Conversion successful!")
#                 return True
#             else:
#                 raise FileNotFoundError(f"Failed to create {pdf_filename}")
#         except Exception as e:
#             st.error(f"Attempt {attempt + 1} failed: {e}")
#             if attempt < retries - 1:
#                 time.sleep(delay)
#             else:
#                 st.error("Failed to convert Word to PDF after multiple attempts.")
#                 return False

# # Function to create a ZIP file containing both the DOCX and PDF files
# def create_zip_file(word_filename, pdf_filename, zip_filename):
#     try:
#         with zipfile.ZipFile(zip_filename, 'w') as zipf:
#             if os.path.exists(word_filename):
#                 zipf.write(word_filename)
#             if os.path.exists(pdf_filename):
#                 zipf.write(pdf_filename)
#         st.success(f"Zip file {zip_filename} created successfully.")
#     except Exception as e:
#         st.error(f"Failed to create zip file: {e}")


# def send_email_with_attachments(subject, body, attachments):
#     to_email = ["hratana261@gmail.com", "khengdalish21@gmail.com", "chlakhna702@gmail.com"]

#     msg = MIMEMultipart()
#     msg['From'] = from_email
#     msg['To'] = ", ".join(to_email)
#     msg['Subject'] = subject

#     msg.attach(MIMEText(body, 'plain'))

#     for attachment in attachments:
#         try:
#             with open(attachment, "rb") as f:
#                 part = MIMEBase('application', 'octet-stream')
#                 part.set_payload(f.read())
#                 encoders.encode_base64(part)
#                 part.add_header('Content-Disposition', f'attachment; filename={os.path.basename(attachment)}')
#                 msg.attach(part)
#         except Exception as e:
#             st.error(f"Failed to attach file {attachment}: {e}")

#     try:
#         server = smtplib.SMTP('smtp.gmail.com', 587)
#         server.starttls()
#         server.login(from_email, password)
#         server.sendmail(from_email, to_email, msg.as_string())
#         server.quit()
#         st.success(f"Email sent to {', '.join(to_email)}")
#     except Exception as e:
#         st.error(f"Failed to send email: {e}")

# def fetch_data(google_sheet_url):
#     try:
#         df = pd.read_csv(google_sheet_url)
#     except Exception as e:
#         st.error(f"Failed to fetch data from Google Sheets: {e}")
#         return None
#     return df

# def fetch_pivot_data(pivot_table_url):
#     try:
#         pivot_df = pd.read_csv(pivot_table_url)
#     except Exception as e:
#         st.error(f"Failed to fetch pivot table data from Google Sheets: {e}")
#         return None
#     return pivot_df

# def ask_about_data(data, question):
#     try:
#         prompt = f"Given the following data:\n{data.to_csv(index=False)}\nAnswer the following question: {question}"

#         response = openai.ChatCompletion.create(
#             model="gpt-4",
#             messages=[{"role": "system", "content": "You are a data analyst."}, {"role": "user", "content": prompt}],
#             max_tokens=500,
#             temperature=0.5
#         )

#         answer = response.choices[0].message['content'].strip()
#         return answer
#     except Exception as e:
#         st.error(f"Failed to get a response: {e}")
#         return None

# def dashboard():
#     st.set_page_config(
#         page_title="DCx Co.,ltd",
#         page_icon="https://dcxsea.com/asset/images/logo/LOGO_DCX.png",
#         layout="wide",
#         initial_sidebar_state="collapsed"
#     )

#     hide_st_style = """
#                 <style>
#                 #MainMenu {visibility: hidden;}
#                 footer {visibility: hidden;}
#                 </style>
#                 """
#     st.markdown(hide_st_style, unsafe_allow_html=True)

#     st.markdown(
#         """
#         <div style="display: flex; align-items: center;">
#             <img src="https://cdn3d.iconscout.com/3d/free/thumb/free-line-chart-growth-3814121-3187502.png" alt="logo" style="width: 90px; margin-right: 15px;">
#             <h3 style="font-family: 'Khmer OS Muol Light', Arial, sans-serif; margin-top: 0;">ការបន្សាំកសិកម្មជនជាតិដើមភាគតិច</h3>
#         </div>
#         """,
#         unsafe_allow_html=True
#     )

#     st.sidebar.markdown(
#         """
#         <div style="display: flex; justify-content: center;margin-top: 0px; margin-bottom: 20px;">
#             <img src="https://dcxsea.com/asset/images/logo/LOGO_DCX.png" style="width: 150px;">
#         </div>
#         """,
#         unsafe_allow_html=True
#     )

#     options = st.sidebar.selectbox(
#         'Choose Dataset',
#         [' ', '6 Months', 'One Year', '6 & 12 Months']
#     )

#     if options == 'One Year':
#         report_title = 'One Year Report'
#         df = fetch_data(
#             google_sheet_url='https://docs.google.com/spreadsheets/d/e/2PACX-1vSnUF27sotZoKCfxKc-dWsLXlKaObixAwluYlygi2GxapQ0QwuFNZkP-3Je_y1YkY8tXgaxm7szHei1/pub?gid=2140672542&single=true&output=csv'
#         )
#     elif options == '6 Months':
#         report_title = '6 Months Report'
#         df = fetch_data(
#             google_sheet_url='https://docs.google.com/spreadsheets/d/e/2PACX-1vSnUF27sotZoKCfxKc-dWsLXlKaObixAwluYlygi2GxapQ0QwuFNZkP-3Je_y1YkY8tXgaxm7szHei1/pub?gid=0&single=true&output=csv'
#         )
#     elif options == '6 & 12 Months':
#         report_title = '6 & 12 Months Report'
#         df = fetch_data(
#             google_sheet_url='https://docs.google.com/spreadsheets/d/e/2PACX-1vSnUF27sotZoKCfxKc-dWsLXlKaObixAwluYlygi2GxapQ0QwuFNZkP-3Je_y1YkY8tXgaxm7szHei1/pub?gid=1666040136&single=true&output=csv'
#         )
#     else:
#         report_title = None
#         df = None
#         st.markdown(
#             """
#             <div style="display: flex; align-items: center;">
#                 <img src="https://symbolshub.org/wp-content/uploads/2019/10/bullet-point-symbol.png" alt="logo" style="width: 25px; margin-right: 5px; vertical-align: middle;">
#                 <h3 style="font-family: 'Khmer OS Muol Light', Arial, sans-serif; margin-top: 0; font-size: 18px; font-weight: bold; vertical-align: middle;">ទិន្នន័យអំពីការបន្សាំកសិកម្មនៃជនជាតិភាគតិច</h3><br><br><br>
#             </div>
#             """,
#             unsafe_allow_html=True
#         )
#         pivot_table_url = 'https://docs.google.com/spreadsheets/d/e/2PACX-1vSnUF27sotZoKCfxKc-dWsLXlKaObixAwluYlygi2GxapQ0QwuFNZkP-3Je_y1YkY8tXgaxm7szHei1/pub?gid=254021688&single=true&output=csv'
#         pivot_df = fetch_pivot_data(pivot_table_url)
#         if pivot_df is not None:
#             pivot = pivot_df.style.set_properties(**{'background-color': 'rgb(161, 219, 255, 0.3)', 'color': 'white'})
#             st.dataframe(pivot)

#     if options in ['One Year', '6 Months', '6 & 12 Months'] and df is not None:
#         st.markdown(
#             """
#             <div style="display: flex; align-items: center;">
#                 <img src="https://symbolshub.org/wp-content/uploads/2019/10/bullet-point-symbol.png" alt="logo" style="width: 25px; margin-right: 5px; vertical-align: middle;">
#                 <h3 style="font-family: 'Khmer OS Muol Light', Arial, sans-serif; margin-top: 0; font-size: 18px; font-weight: bold; vertical-align: middle;">ទិន្នន័យអំពីការបន្សាំកសិកម្មនៃជនជាតិភាគតិច</h3><br><br><br>
#             </div>
#             """,
#             unsafe_allow_html=True
#         )
#         df_style = df.style.set_properties(**{'background-color': 'rgb(161, 219, 255, 0.3)', 'color': 'white'})
#         st.dataframe(df_style)

#         st.markdown(
#             """
#             <div style="display: flex; align-items: center;">
#                 <label for="question_input" style="font-family: 'Khmer OS Muol Light', Arial, sans-serif; font-size: 14px;">សូមបញ្ចូលសំណួរ:</label>
#             </div>
#             """,
#             unsafe_allow_html=True
#         )

#         col1, col2 = st.columns([3, 1])
#         with col1:
#             # Text input field
#             question = st.text_input(" ", key="question_input_key")

#         with col2:
#             st.markdown("""
#                 <style>
#                 div.stButton > button {
#                     width: auto;
#                     height: 40px;
#                     margin-top: 11px;
#                 }
#                 div.stButton > button:hover {
#                     background-color: darkgreen;
#                 }
#                 </style>
#                 """, unsafe_allow_html=True)
#             search = st.button("Search")

#         # Display the answer if the button is clicked
#         if search and question:
#             st.session_state["search_result"] = ask_about_data(df, question)

#         if "search_result" in st.session_state:
#             st.write(st.session_state["search_result"])

#          # Add button to generate report with custom style
#         button_style = """
#             <style>
#             .stButton>button {
#                 background-color: green;
#                 margin-buttom: 24px;
#                 color: white;
#                 border: none;
#                 padding: 0.5em 1em;
#                 cursor: pointer;
#             }
#             .stButton>button:hover {
#                 background-color: darkgreen;
#                 color: #DFFF00;
#             }
#             </style>
#         """
#         st.markdown(button_style, unsafe_allow_html=True)

#         if st.button('Generate Report'):
#             if df.empty:
#                 st.error("No data available to send.")
#             else:
#                 df_cleaned = df.fillna('')  # Fill NaNs with empty strings
#                 data = df_cleaned.to_dict(orient='records')  # Convert DataFrame to dictionary format

#                 # Generate report with the ChatGPT API
#                 if "report_parts" not in st.session_state:
#                     report_content, word_filename, pdf_filename = generate_report_with_chatgpt(data, report_title)
#                 else:
#                     report_content = st.session_state["report_parts"]
#                     word_filename = f'{report_title}.docx'
#                     pdf_filename = f'{report_title}.pdf'
#                     save_report_as_word(report_content, word_filename)
#                     convert_to_pdf_with_retry(word_filename, pdf_filename)

#                 if report_content:
#                     zip_filename = f'{report_title}.zip'
#                     create_zip_file(word_filename, pdf_filename, zip_filename)

#                     # Send email with the zip file
#                     send_email_with_attachments(f"{report_title} Generated Report", "Please find the attached reports.", [zip_filename])

#                     # Send report to Telegram
#                     send_to_telegram(word_filename, f"Here is your generated {report_title} (Word).")
#                     send_to_telegram(pdf_filename, f"Here is your generated {report_title} (PDF).")

#                     # Custom style for the download button
#                     st.markdown("""
#                     <style>
#                     div.stDownloadButton > button {
#                         background-color: #006400;
#                         color: white;
#                         padding: 10px;
#                         font-size: 16px;
#                         border-radius: 5px;
#                         border: none;
#                         cursor: pointer;
#                     }
#                     div.stDownloadButton > button:hover {
#                         background-color: #005000;
#                         color: #DFFF00;
#                     }
#                     </style>
#                     """, unsafe_allow_html=True)

#                     # Allow downloading of the zip file
#                     st.download_button(f'Download {report_title}', data=open(zip_filename, 'rb').read(), file_name=zip_filename, mime='application/zip')
#                 else:
#                     st.write("Failed to generate report.")

# if __name__ == '__main__':
#     dashboard()





import streamlit as st
import pandas as pd
import json
import os
import zipfile
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import openai
from datetime import datetime
import re
import requests
from io import BytesIO
from dotenv import load_dotenv
import time
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
from reportlab.lib.units import inch
from reportlab.lib import colors

# Load environment variables from .env file
load_dotenv()

# Set your OpenAI API key from environment variable
openai.api_key = st.secrets["OPENAI_API_KEY"]

# Set your Telegram bot token and chat ID
telegram_bot_token = os.getenv('TELEGRAM_BOT_TOKEN')
telegram_chat_id = '-1002164741954'

# Email configuration
from_email = "seaklav168@gmail.com"
password = os.getenv('EMAIL_PASSWORD')

def send_to_telegram(file_path, caption):
    if not os.path.exists(file_path):
        st.error(f"File {file_path} not found. Cannot send to Telegram.")
        return
    url = f"https://api.telegram.org/bot{telegram_bot_token}/sendDocument"
    with open(file_path, 'rb') as file:
        response = requests.post(url, data={'chat_id': telegram_chat_id, 'caption': caption}, files={'document': file})
    return response

def generate_report_with_chatgpt(data, report_title):
    try:
        prompt = f"Please generate a formal report based on the following data: {json.dumps(data, indent=2)}"
        
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[{"role": "system", "content": "You are a helpful assistant."}, {"role": "user", "content": prompt}],
            max_tokens=8000,
            temperature=0.7
        )

        report_parts = response.choices[0].message['content'].strip()

        st.write(report_parts)

        # Save report as Word and PDF documents
        word_filename = f'{report_title}.docx'
        pdf_filename = f'{report_title}.pdf'
        save_report_as_word(report_parts, word_filename)
        convert_to_pdf_with_retry(word_filename, pdf_filename)

        return report_parts, word_filename, pdf_filename
    except Exception as e:
        st.error(f"Failed to generate report: {e}")
        return None, None, None

def create_cover_page(doc, report_title):
    section = doc.sections[0]
    section.page_height = Pt(842)  # A4 size height
    section.page_width = Pt(595)   # A4 size width

    # Set margins to 2 cm (56.7 points)
    for section in doc.sections:
        section.top_margin = Pt(56.7)
        section.bottom_margin = Pt(56.7)
        section.left_margin = Pt(56.7)
        section.right_margin = Pt(56.7)

    # Add logo from URL to the cover page
    logo_url = "https://dcxsea.com/asset/images/logo/LOGO_DCX.png"
    response = requests.get(logo_url)
    if response.status_code == 200:
        logo_image = BytesIO(response.content)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        run.add_picture(logo_image, width=Pt(150))  # Adjust width as needed

        # Add company name after the logo
        run = paragraph.add_run("  DCx Co., Ltd.")
        run.font.size = Pt(24)
        run.bold = True
    else:
        st.error("Failed to download the logo image.")

    doc.add_paragraph("\n")
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = cover.add_run("Report")
    run.font.size = Pt(20)
    run.bold = True
    cover.add_run("\n").font.size = Pt(24)

    doc.add_paragraph("\n")
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = cover.add_run("Indigenous Agriculture Adaptation")
    run.font.size = Pt(24)
    run.bold = True
    cover.add_run("\n").font.size = Pt(24)

    for _ in range(2):
        doc.add_paragraph("\n")

    # Add prepared for
    prepared_for = doc.add_paragraph()
    prepared_for.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prepared_for.add_run("Prepared for: Jack Jasmin")
    run.font.size = Pt(14)
    run.add_break()  # Add a break line
    run.add_break()
    run = prepared_for.add_run("Prepared by: Black Eye Team")
    run.font.size = Pt(14)
    run.add_break()
    run.add_break()

    # Place the date at the bottom center of the cover page
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_paragraph.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(14)

    # Create a new section for the rest of the document, so the footer with the date is not repeated
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Ensure the new section starts with a fresh header and footer
    new_section.footer.is_linked_to_previous = False
    new_section.footer.paragraphs[0].clear()
    new_section.header.is_linked_to_previous = False

def add_markdown_formatted_text(paragraph, text):
    bold_italic_pattern = re.compile(r'\*\*\*(.+?)\*\*\*')
    bold_pattern = re.compile(r'\*\*(.+?)\*\*')
    italic_pattern = re.compile(r'\*(.+?)\*')
    patterns = [
        ('bold_italic', bold_italic_pattern),
        ('bold', bold_pattern),
        ('italic', italic_pattern)
    ]

    def replace_match(match, style):
        if style == 'bold_italic':
            run = paragraph.add_run(match.group(1))
            run.bold = True
            run.italic = True
        elif style == 'bold':
            run = paragraph.add_run(match.group(1))
            run.bold = True
        elif style == 'italic':
            run = paragraph.add_run(match.group(1))
            run.italic = True

    cursor = 0
    text_length = len(text)
    while cursor < text_length:
        nearest_match = None
        nearest_style = None
        nearest_start = text_length

        for style, pattern in patterns:
            match = pattern.search(text, cursor)
            if match:
                start = match.start()
                if start < nearest_start:
                    nearest_match = match
                    nearest_style = style
                    nearest_start = start

        if nearest_match:
            if nearest_start > cursor:
                paragraph.add_run(text[cursor:nearest_start])
            replace_match(nearest_match, nearest_style)
            cursor = nearest_match.end()
        else:
            paragraph.add_run(text[cursor:])
            break

    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def save_report_as_word(report, filename):
    try:
        doc = Document()

        for section in doc.sections:
            section.top_margin = Pt(56.7)
            section.bottom_margin = Pt(56.7)
            section.left_margin = Pt(56.7)
            section.right_margin = Pt(56.7)

        create_cover_page(doc, filename.split('.')[0])

        lines = report.split('\n')
        table = None
        lines = lines[1:]

        for line in lines:
            if line.strip() == "---":
                doc.add_page_break()
            elif line.strip().startswith("# "):
                paragraph = doc.add_heading(line.strip()[2:], level=1)
            elif line.strip().startswith("## "):
                paragraph = doc.add_heading(line.strip()[3:], level=2)
            elif line.strip().startswith("### "):
                paragraph = doc.add_heading(line.strip()[4:], level=3)
            elif line.strip().startswith("#### "):
                paragraph = doc.add_heading(line.strip()[5:], level=4)
            elif line.strip().startswith("* "):
                doc.add_paragraph(line.strip()[2:], style='List Bullet')
            elif "|" in line:
                table_data = [cell.strip() for cell in line.split('|') if cell]
                if not table:
                    table = doc.add_table(rows=1, cols=len(table_data))
                    hdr_cells = table.rows[0].cells
                    for i, cell_data in enumerate(table_data):
                        hdr_cells[i].text = cell_data
                else:
                    row_cells = table.add_row().cells
                    for i, cell_data in enumerate(table_data):
                        row_cells[i].text = cell_data
            else:
                paragraph = doc.add_paragraph()
                add_markdown_formatted_text(paragraph, line)

        doc.save(filename)
    except Exception as e:
        st.error(f"Failed to save Word report: {e}")

def read_docx_content(word_filename):
    doc = Document(word_filename)
    content = ""
    for para in doc.paragraphs:
        content += para.text + "\n"
    return content

def convert_to_pdf_with_retry(word_filename, pdf_filename, retries=3, delay=5):
    for attempt in range(retries):
        try:
            content = read_docx_content(word_filename)
            create_pdf_with_structure(pdf_filename, content)
            if os.path.exists(pdf_filename):
                st.success("Conversion successful!")
                return True
            else:
                raise FileNotFoundError(f"Failed to create {pdf_filename}")
        except Exception as e:
            st.error(f"Attempt {attempt + 1} failed: {e}")
            if attempt < retries - 1:
                time.sleep(delay)
            else:
                st.error("Failed to convert Word to PDF after multiple attempts.")
                return False

def create_pdf_with_structure(pdf_filename, content):
    try:
        doc = SimpleDocTemplate(pdf_filename, pagesize=A4,
                                rightMargin=56.7, leftMargin=56.7,
                                topMargin=56.7, bottomMargin=56.7)
        styles = getSampleStyleSheet()

        # Use existing styles or create new custom styles
        heading1_style = styles['Heading1']
        heading2_style = styles['Heading2']
        heading3_style = styles['Heading3']
        body_text_style = styles['BodyText']
        justified_style = ParagraphStyle(name="Justified", parent=styles['BodyText'], alignment=4)

        elements = []

        # Add cover page
        logo_url = "https://dcxsea.com/asset/images/logo/LOGO_DCX.png"
        response = requests.get(logo_url)
        if response.status_code == 200:
            logo_image = BytesIO(response.content)
            img = Image(logo_image, 2*inch, 1*inch)
            img.hAlign = 'CENTER'
            elements.append(img)

        elements.append(Spacer(1, 12))
        elements.append(Paragraph("<b>DCx Co., Ltd.</b>", heading1_style))
        elements.append(Spacer(1, 24))
        elements.append(Paragraph("<b>Report</b>", heading1_style))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph("Indigenous Agriculture Adaptation", heading1_style))
        elements.append(Spacer(1, 48))
        elements.append(Paragraph("Prepared for: Jack Jasmin", body_text_style))
        elements.append(Paragraph("Prepared by: Black Eye Team", body_text_style))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", body_text_style))
        elements.append(PageBreak())

        # Process content into structured paragraphs
        lines = content.split("\n")
        for line in lines:
            if line.strip().startswith("# "):
                elements.append(Paragraph(line.strip()[2:], heading1_style))
            elif line.strip().startswith("## "):
                elements.append(Paragraph(line.strip()[3:], heading2_style))
            elif line.strip().startswith("### "):
                elements.append(Paragraph(line.strip()[4:], heading3_style))
            elif line.strip().startswith("* "):
                elements.append(Paragraph(f"• {line.strip()[2:]}", body_text_style))
            else:
                elements.append(Paragraph(line.strip(), justified_style))
            elements.append(Spacer(1, 12))

        doc.build(elements)

        if not os.path.exists(pdf_filename):
            raise FileNotFoundError(f"{pdf_filename} not created.")
    except Exception as e:
        st.error(f"Failed to create PDF: {e}")
        raise RuntimeError(f"Failed to create PDF: {e}")

def send_email_with_attachments(subject, body, attachments):
    to_email = ["hratana261@gmail.com", "khengdalish21@gmail.com", "chlakhna702@gmail.com"]

    msg = MIMEMultipart()
    msg['From'] = from_email
    msg['To'] = ", ".join(to_email)
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'plain'))

    for attachment in attachments:
        try:
            with open(attachment, "rb") as f:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(f.read())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', f'attachment; filename={os.path.basename(attachment)}')
                msg.attach(part)
        except Exception as e:
            st.error(f"Failed to attach file {attachment}: {e}")

    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(from_email, password)
        server.sendmail(from_email, to_email, msg.as_string())
        server.quit()
        st.success(f"Email sent to {', '.join(to_email)}")
    except Exception as e:
        st.error(f"Failed to send email: {e}")

def fetch_data(google_sheet_url):
    try:
        df = pd.read_csv(google_sheet_url)
    except Exception as e:
        st.error(f"Failed to fetch data from Google Sheets: {e}")
        return None
    return df

def fetch_pivot_data(pivot_table_url):
    try:
        pivot_df = pd.read_csv(pivot_table_url)
    except Exception as e:
        st.error(f"Failed to fetch pivot table data from Google Sheets: {e}")
        return None
    return pivot_df

def ask_about_data(data, question):
    try:
        prompt = f"Given the following data:\n{data.to_csv(index=False)}\nAnswer the following question: {question}"

        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[{"role": "system", "content": "You are a data analyst."}, {"role": "user", "content": prompt}],
            max_tokens=500,
            temperature=0.5
        )

        answer = response.choices[0].message['content'].strip()
        return answer
    except Exception as e:
        st.error(f"Failed to get a response: {e}")
        return None

def dashboard():
    st.set_page_config(
        page_title="DCx Co.,ltd",
        page_icon="https://dcxsea.com/asset/images/logo/LOGO_DCX.png",
        layout="wide",
        initial_sidebar_state="collapsed"
    )

    hide_st_style = """
                <style>
                #MainMenu {visibility: hidden;}
                footer {visibility: hidden;}
                </style>
                """
    st.markdown(hide_st_style, unsafe_allow_html=True)

    st.markdown(
        """
        <div style="display: flex; align-items: center;">
            <img src="https://cdn3d.iconscout.com/3d/free/thumb/free-line-chart-growth-3814121-3187502.png" alt="logo" style="width: 90px; margin-right: 15px;">
            <h3 style="font-family: 'Khmer OS Muol Light', Arial, sans-serif; margin-top: 0;">ការបន្សាំកសិកម្មជនជាតិដើមភាគតិច</h3>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.sidebar.markdown(
        """
        <div style="display: flex; justify-content: center;margin-top: 0px; margin-bottom: 20px;">
            <img src="https://dcxsea.com/asset/images/logo/LOGO_DCX.png" style="width: 150px;">
        </div>
        """,
        unsafe_allow_html=True
    )

    options = st.sidebar.selectbox(
        'Choose Dataset',
        [' ', '6 Months', 'One Year', '6 & 12 Months']
    )

    if options == 'One Year':
        report_title = 'One Year Report'
        df = fetch_data(
            google_sheet_url='https://docs.google.com/spreadsheets/d/e/2PACX-1vSnUF27sotZoKCfxKc-dWsLXlKaObixAwluYlygi2GxapQ0QwuFNZkP-3Je_y1YkY8tXgaxm7szHei1/pub?gid=2140672542&single=true&output=csv'
        )
    elif options == '6 Months':
        report_title = '6 Months Report'
        df = fetch_data(
            google_sheet_url='https://docs.google.com/spreadsheets/d/e/2PACX-1vSnUF27sotZoKCfxKc-dWsLXlKaObixAwluYlygi2GxapQ0QwuFNZkP-3Je_y1YkY8tXgaxm7szHei1/pub?gid=0&single=true&output=csv'
        )
    elif options == '6 & 12 Months':
        report_title = '6 & 12 Months Report'
        df = fetch_data(
            google_sheet_url='https://docs.google.com/spreadsheets/d/e/2PACX-1vSnUF27sotZoKCfxKc-dWsLXlKaObixAwluYlygi2GxapQ0QwuFNZk-3Je_y1YkY8tXgaxm7szHei1/pub?gid=1666040136&single=true&output=csv'
        )
    else:
        report_title = None
        df = None
        st.markdown(
            """
            <div style="display: flex; align-items: center;">
                <img src="https://symbolshub.org/wp-content/uploads/2019/10/bullet-point-symbol.png" alt="logo" style="width: 25px; margin-right: 5px; vertical-align: middle;">
                <h3 style="font-family: 'Khmer OS Muol Light', Arial, sans-serif; margin-top: 0; font-size: 18px; font-weight: bold; vertical-align: middle;">ទិន្នន័យអំពីការបន្សាំកសិកម្មនៃជនជាតិភាគតិច</h3><br><br><br>
            </div>
            """,
            unsafe_allow_html=True
        )
        pivot_table_url = 'https://docs.google.com/spreadsheets/d/e/2PACX-1vSnUF27sotZoKCfxKc-dWsLXlKaObixAwluYlygi2GxapQ0QwuFNZk-3Je_y1YkY8tXgaxm7szHei1/pub?gid=254021688&single=true&output=csv'
        pivot_df = fetch_pivot_data(pivot_table_url)
        if pivot_df is not None:
            pivot = pivot_df.style.set_properties(**{'background-color': 'rgb(161, 219, 255, 0.3)', 'color': 'white'})
            st.dataframe(pivot)

    if options in ['One Year', '6 Months', '6 & 12 Months'] and df is not None:
        st.markdown(
            """
            <div style="display: flex; align-items: center;">
                <img src="https://symbolshub.org/wp-content/uploads/2019/10/bullet-point-symbol.png" alt="logo" style="width: 25px; margin-right: 5px; vertical-align: middle;">
                <h3 style="font-family: 'Khmer OS Muol Light', Arial, sans-serif; margin-top: 0; font-size: 18px; font-weight: bold; vertical-align: middle;">ទិន្នន័យអំពីការបន្សាំកសិកម្មនៃជនជាតិភាគតិច</h3><br><br><br>
            </div>
            """,
            unsafe_allow_html=True
        )
        df_style = df.style.set_properties(**{'background-color': 'rgb(161, 219, 255, 0.3)', 'color': 'white'})
        st.dataframe(df_style)

        st.markdown(
            """
            <div style="display: flex; align-items: center;">
                <label for="question_input" style="font-family: 'Khmer OS Muol Light', Arial, sans-serif; font-size: 14px;">សូមបញ្ចូលសំណួរ:</label>
            </div>
            """,
            unsafe_allow_html=True
        )

        col1, col2 = st.columns([3, 1])
        with col1:
            # Text input field
            question = st.text_input(" ", key="question_input_key")

        with col2:
            st.markdown("""
                <style>
                div.stButton > button {
                    width: auto;
                    height: 40px;
                    margin-top: 11px;
                }
                div.stButton > button:hover {
                    background-color: darkgreen;
                }
                </style>
                """, unsafe_allow_html=True)
            search = st.button("Search")

        # Display the answer if the button is clicked
        if search and question:
            st.session_state["search_result"] = ask_about_data(df, question)

        if "search_result" in st.session_state:
            st.write(st.session_state["search_result"])

         # Add button to generate report with custom style
        button_style = """
            <style>
            .stButton>button {
                background-color: green;
                margin-buttom: 24px;
                color: white;
                border: none;
                padding: 0.5em 1em;
                cursor: pointer;
            }
            .stButton>button:hover {
                background-color: darkgreen;
                color: #DFFF00;
            }
            </style>
        """
        st.markdown(button_style, unsafe_allow_html=True)

        if st.button('Generate Report'):
            if df.empty:
                st.error("No data available to send.")
            else:
                df_cleaned = df.fillna('')  # Fill NaNs with empty strings
                data = df_cleaned.to_dict(orient='records')  # Convert DataFrame to dictionary format

                # Generate report with the ChatGPT API
                if "report_parts" not in st.session_state:
                    report_content, word_filename, pdf_filename = generate_report_with_chatgpt(data, report_title)
                else:
                    report_content = st.session_state["report_parts"]
                    word_filename = f'{report_title}.docx'
                    pdf_filename = f'{report_title}.pdf'
                    save_report_as_word(report_content, word_filename)
                    convert_to_pdf_with_retry(word_filename, pdf_filename)

                if report_content:
                    zip_filename = f'{report_title}.zip'
                    create_zip_file(word_filename, pdf_filename, zip_filename)

                    # Send email with the zip file
                    send_email_with_attachments(f"{report_title} Generated Report", "Please find the attached reports.", [zip_filename])

                    # Send report to Telegram
                    send_to_telegram(word_filename, f"Here is your generated {report_title} (Word).")
                    send_to_telegram(pdf_filename, f"Here is your generated {report_title} (PDF).")

                    # Custom style for the download button
                    st.markdown("""
                    <style>
                    div.stDownloadButton > button {
                        background-color: #006400;
                        color: white;
                        padding: 10px;
                        font-size: 16px;
                        border-radius: 5px;
                        border: none;
                        cursor: pointer;
                    }
                    div.stDownloadButton > button:hover {
                        background-color: #005000;
                        color: #DFFF00;
                    }
                    </style>
                    """, unsafe_allow_html=True)

                    # Allow downloading of the zip file
                    st.download_button(f'Download {report_title}', data=open(zip_filename, 'rb').read(), file_name=zip_filename, mime='application/zip')
                else:
                    st.write("Failed to generate report.")

if __name__ == '__main__':
    dashboard()
